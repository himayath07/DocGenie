import streamlit as st
import os
import pandas as pd
import tempfile
import asyncio
import shutil
from concurrent.futures import ThreadPoolExecutor
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings import HuggingFaceEmbeddings
from langchain import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import Document
import docx
import speech_recognition as sr
from gtts import gTTS

import io
import pygame
from transformers import pipeline
import pytesseract
from PIL import Image
from transformers import pipeline
from youtube_transcript_api import YouTubeTranscriptApi
import re
from langchain.prompts import PromptTemplate
import graphviz
from langchain.llms import HuggingFaceHub
from huggingface_hub import login, HfApi
import requests
from fastapi import FastAPI, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.cors import CORSMiddleware
from starlette.responses import JSONResponse
import uvicorn
from fastapi.staticfiles import StaticFiles

os.environ["HUGGINGFACE_API_TOKEN"] = "your_huggingface_token_here"  # Replace with your actual token

pygame.mixer.init()

@st.cache_resource
def get_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

def validate_hf_token(token):
    try:
        api = HfApi(token=token)
        api.whoami()
        return True
    except Exception:
        return False

def generate_mindmap(text, map_type="mindmap"):
    """
    Generate a Mermaid diagram (mindmap/flowchart) from text
    """
    try:
        # Verify token is available
        if "HUGGINGFACEHUB_API_TOKEN" not in st.session_state:
            return "Please set your HuggingFace API token first."
        
        # Create template based on map_type
        if map_type == "mindmap":
            template = """
            Create a mindmap in Mermaid syntax based on the following text. 
            Focus on the main concepts and their relationships.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the mindmap like this:
            mindmap
                Root Topic
                    Subtopic 1
                        Detail A
                        Detail B
                    Subtopic 2
                        Detail C
                        Detail D
            """
        elif map_type == "flowchart":
            template = """
            Create a flowchart in Mermaid syntax based on the following text.
            Focus on the process flow and relationships between concepts.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the flowchart like this:
            graph TD
                A[Start] --> B[Process 1]
                B --> C[Process 2]
                C --> D[End]
            """
        else:  # graph
            template = """
            Create a graph diagram in Mermaid syntax based on the following text.
            Focus on the relationships between concepts.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the graph like this:
            graph LR
                A[Concept 1] --- B[Concept 2]
                B --- C[Concept 3]
                C --- D[Concept 4]
            """
        
        prompt = PromptTemplate(
            input_variables=["text"],
            template=template
        )
        
        # Initialize the model with explicit token
        llm = HuggingFaceHub(
            repo_id="mistralai/Mistral-7B-Instruct-v0.1",
            model_kwargs={
                "temperature": 0.7,
                "max_new_tokens": 512,
                "top_p": 0.95,
                "repetition_penalty": 1.15
            },
            huggingfacehub_api_token=st.session_state["HUGGINGFACEHUB_API_TOKEN"]
        )
        
        # Generate the diagram
        diagram_code = llm.predict(prompt.format(text=text))
        return diagram_code.strip()
    
    except requests.exceptions.HTTPError as e:
        st.error("Authentication error. Please check your HuggingFace token.")
        return f"Error: Authentication failed. Please check your token."
    except Exception as e:
        st.error(f"Error details: {str(e)}")
        return f"Error generating diagram: {str(e)}"

def summarize_document(file_path, file_type):
    """
    Summarizes the content of the uploaded document.
    This function loads the document using your existing loader logic,
    then summarizes the combined text from the document pages.
    """
    # Example: For PDF/TXT/DOCX you can use your existing create_document_db logic
    # Here we'll simply open a TXT file for demonstration purposes.
    # You can adjust it based on your document loader (for PDF, DOCX, etc.)
    try:
        if file_type == 'pdf':
            from langchain.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            pages = loader.load_and_split()
            text = "\n".join([page.page_content for page in pages])
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_type in ['word', 'doc']:
            import docx
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            text = "Unsupported file type for summarization."
    except Exception as e:
        st.error(f"Error while processing file: {e}")
        return None

    summarizer = get_summarizer()
    summary = summarizer(text, max_length=150, min_length=50, do_sample=False)
    return summary[0]['summary_text']

st.set_page_config(
    page_title="Document Interaction Platform",
    page_icon="📄",
    layout="wide"
)

# Add custom CSS
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton button {
        width: 100%;
    }
    .upload-section {
        padding: 2rem;
        border-radius: 0.5rem;
        background-color: #f0f2f6;
        margin-bottom: 2rem;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📄 Document Interaction Platform")
st.markdown("""
    Welcome to the Document Interaction Platform! This tool allows you to:
    - 📝 Upload and summarize documents
    - 🔄 Convert between different file formats
    - 📊 Process various file types including PDF, Word, Text, and Images
""")

# Add this right after your imports, before any Streamlit UI code
def process_query(file, query):
    """
    Process a query about the document using LangChain
    Args:
        file: The uploaded file object
        query: The user's question about the document
    Returns:
        str: Response to the query
    """
    try:
        # Read the content of the file
        if file.type == "application/pdf":
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        elif file.type == "text/plain":
            text = file.getvalue().decode()
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            import docx
            doc = docx.Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        else:
            return "Unsupported file format. Please upload a PDF, TXT, or DOCX file."

        # Create a simple response (replace this with more sophisticated processing)
        from langchain.text_splitter import CharacterTextSplitter
        from langchain.embeddings import HuggingFaceEmbeddings
        from langchain.vectorstores import FAISS
        from langchain.chains.question_answering import load_qa_chain
        from langchain.llms import HuggingFaceHub

        # Split the text into chunks
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)

        # Create embeddings
        embeddings = HuggingFaceEmbeddings()
        knowledge_base = FAISS.from_texts(chunks, embeddings)

        # Get relevant chunks
        docs = knowledge_base.similarity_search(query)

        # Create chain
        chain = load_qa_chain(HuggingFaceHub(repo_id="google/flan-t5-large"), chain_type="stuff")
        
        # Get response
        response = chain.run(input_documents=docs, question=query)
        return response

    except Exception as e:
        return f"Error processing query: {str(e)}"

# File upload section in a container
with st.container():
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.subheader("📁 Upload Your Document")
    
    try:
        # Add debug message
        st.info("Attempting file upload...")
        
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['pdf', 'txt', 'docx', 'doc', 'jpg', 'jpeg', 'png'],
            help="Supported formats: PDF, TXT, DOCX/DOC, JPG, PNG"
        )
        
        if uploaded_file is not None:
            # Display file details for debugging
            st.write(f"File details: {uploaded_file.name}, Size: {uploaded_file.size} bytes, Type: {uploaded_file.type}")
            
            # Validate file size
            if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit
                st.error("File size exceeds 10MB limit. Please upload a smaller file.")
                uploaded_file = None
            else:
                try:
                    # Create temp directory if it doesn't exist
                    temp_dir = tempfile.gettempdir()
                    if not os.path.exists(temp_dir):
                        os.makedirs(temp_dir)
                    
                    # Save the uploaded file to a temporary path with error handling
                    file_path = os.path.join(temp_dir, uploaded_file.name)
                    
                    # Debug message
                    st.info(f"Attempting to save file to: {file_path}")
                    
                    with open(file_path, 'wb') as f:
                        f.write(uploaded_file.getbuffer())
                        
                    if os.path.exists(file_path):
                        st.success(f"📄 Successfully uploaded: {uploaded_file.name}")
                    else:
                        st.error("File was not saved properly")
                        
                except Exception as e:
                    st.error(f"Error saving file: {str(e)}")
                    st.info("Please try uploading again or contact support if the issue persists.")
                    st.write("Error details:", str(e))  # More detailed error information
                    uploaded_file = None
                    
    except Exception as e:
        st.error(f"Upload error: {str(e)}")
        st.info("Please check your internet connection and try again.")
        st.write("Error details:", str(e))  # More detailed error information
        uploaded_file = None
        
    st.markdown('</div>', unsafe_allow_html=True)

# Add this after the file upload section and before the columns
with st.container():
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.subheader("🎤 Voice & Text Query")
    
    # Text input for queries
    text_query = st.text_input(
        "Enter your question about the document",
        placeholder="e.g., What is the main topic of this document?",
        help="Type your question here"
    )
    
    # Voice input
    if st.button("🎤 Start Voice Input"):
        with st.spinner("Listening..."):
            r = sr.Recognizer()
            try:
                with sr.Microphone() as source:
                    st.info("Listening... Speak now!")
                    audio = r.listen(source, timeout=5)
                    query = r.recognize_google(audio)
                    st.session_state['voice_query'] = query
                    st.success(f"Recognized: {query}")
                    text_query = query
            except Exception as e:
                st.error(f"Sorry, there was an error: {str(e)}")
    
    # Process query (either voice or text)
    if text_query and st.button("🔍 Get Answer"):
        with st.spinner("Processing your query..."):
            try:
                # Add your query processing logic here
                response = process_query(uploaded_file, text_query)
                
                # Display text response
                st.markdown("### 📝 Answer:")
                st.write(response)
                
                # Convert response to speech
                if st.button("🔊 Listen to Response"):
                    with st.spinner("Converting to speech..."):
                        # Create temporary file for audio
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as fp:
                            tts = gTTS(text=response, lang='en')
                            tts.save(fp.name)
                            
                            # Initialize pygame mixer
                            pygame.mixer.init()
                            pygame.mixer.music.load(fp.name)
                            pygame.mixer.music.play()
                            
                            # Wait for audio to finish
                            while pygame.mixer.music.get_busy():
                                pygame.time.Clock().tick(10)
                            
                            # Cleanup
                            pygame.mixer.quit()
                            os.unlink(fp.name)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
    st.markdown('</div>', unsafe_allow_html=True)

file_type = None

def convert_document(file_path, source_type, target_type):
    """
    Converts a file from source_type to target_type.
    Supported source/target types: 
      - "pdf", "word" (DOCX), "txt", "doc", "image" (JPG)
    """
    ext_map = {
        "pdf": ".pdf",
        "word": ".docx",
        "txt": ".txt",
        "doc": ".doc",
        "image": ".jpg"
    }
    if target_type not in ext_map:
        st.error("Unsupported target format.")
        return None
    target_ext = ext_map[target_type]
    base, _ = os.path.splitext(file_path)
    target_file_path = base + f"_converted{target_ext}"
    
    # --- PDF Conversions ---
    if source_type == "pdf" and target_type == "word":
        try:
            from pdf2docx import parse
        except ImportError:
            st.error("pdf2docx is not installed. Please add it to your requirements.")
            return None
        parse(file_path, target_file_path, start=0, end=None)
        return target_file_path

    elif source_type == "pdf" and target_type == "txt":
        from langchain.document_loaders import PyPDFLoader
        loader = PyPDFLoader(file_path)
        pages = loader.load_and_split()
        text = "\n".join([page.page_content for page in pages])
        with open(target_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return target_file_path

    elif source_type == "pdf" and target_type == "image":
        try:
            from pdf2image import convert_from_path
        except ImportError:
            st.error("pdf2image is not installed. Please add it to your requirements.")
            return None
        images = convert_from_path(file_path)
        if images:
            target_file_path = base + "_page1.jpg"
            images[0].save(target_file_path, "JPEG")
            return target_file_path
        else:
            return None

    elif source_type == "pdf" and target_type == "doc":
        converted_word = convert_document(file_path, "pdf", "word")
        if converted_word:
            new_target = os.path.splitext(converted_word)[0] + ".doc"
            os.rename(converted_word, new_target)
            return new_target
        else:
            return None

    # --- Word (or DOC) to PDF ---
    elif source_type in ["word", "doc"] and target_type == "pdf":
        try:
            from docx2pdf import convert as docx2pdf_convert
        except ImportError:
            st.error("docx2pdf is not installed. Please add it to your requirements.")
            return None
        target_file_path = base + "_converted.pdf"
        docx2pdf_convert(file_path, target_file_path)
        return target_file_path

    # --- TXT to PDF ---
    elif source_type == "txt" and target_type == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            st.error("fpdf is not installed. Please add it to your requirements.")
            return None
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        pdf.multi_cell(0, 10, text)
        pdf.output(target_file_path)
        return target_file_path

    # --- Word (or DOC) to TXT ---
    elif source_type in ["word", "doc"] and target_type == "txt":
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        with open(target_file_path, "w", encoding="utf-8") as f:
            f.write(text)
        return target_file_path

    # --- TXT to Word (or DOC) ---
    elif source_type == "txt" and target_type in ["word", "doc"]:
        from docx import Document
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        doc = Document()
        doc.add_paragraph(text)
        doc.save(target_file_path)
        return target_file_path

    # --- Image to PDF ---
    elif source_type == "image" and target_type == "pdf":
        from PIL import Image
        img = Image.open(file_path)
        target_file_path = base + "_converted.pdf"
        img.save(target_file_path, "PDF", resolution=100.0)
        return target_file_path

    else:
        st.error(f"Conversion from {source_type} to {target_type} is not supported.")
        return None


if uploaded_file is not None:
    # Determine file type based on extension
    filename = uploaded_file.name.lower()
    if filename.endswith('.pdf'):
        file_type = 'pdf'
    elif filename.endswith('.txt'):
        file_type = 'txt'
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        file_type = 'word'
    elif filename.endswith('.jpg') or filename.endswith('.jpeg') or filename.endswith('.png'):
        file_type = 'image'
    else:
        st.error("Unsupported file type.")
    
    # Save the uploaded file to a temporary path
    file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
    with open(file_path, 'wb') as f:
        f.write(uploaded_file.getbuffer())
    
    st.success(f"📄 Uploaded file: {uploaded_file.name}")

    # Create columns for better layout
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("### 📝 Summarize Document")
        if st.button("📋 Generate Summary"):
            with st.spinner('Generating summary...'):
                summary = summarize_document(file_path, file_type)
            if summary:
                st.success("Summary generated successfully!")
                st.markdown("**Document Summary:**")
                st.markdown(f">{summary}")
                # Store summary in session state
                st.session_state['document_summary'] = summary
            else:
                st.error("Summarization failed. Please check your file and try again.")
        
        # Add visualization options for document
        if 'document_summary' in st.session_state:
            st.write("### 📊 Document Visualization")
            viz_type = st.selectbox(
                "Select visualization type",
                options=["mindmap", "flowchart", "graph"],
                key="doc_viz_type",
                help="Choose how you want to visualize the document content"
            )
            
            if st.button("Generate Document Visualization"):
                with st.spinner("Generating visualization..."):
                    try:
                        diagram_code = generate_mindmap(st.session_state['document_summary'], viz_type)
                        st.markdown("### 📊 Document Diagram")
                        st.markdown(f"```mermaid\n{diagram_code}\n```")
                        
                        # Add download button for diagram
                        st.download_button(
                            "⬇️ Download Diagram Code",
                            diagram_code,
                            file_name=f"document_{viz_type}_diagram.mmd",
                            mime="text/plain",
                            key="doc_diagram_download"
                        )
                        
                        # Show raw code option
                        if st.checkbox("Show Raw Diagram Code", key="doc_show_code"):
                            st.code(diagram_code, language="mermaid")
                    except Exception as e:
                        st.error(f"Error generating diagram: {str(e)}")
    
    with col2:
        st.write("### 🔄 Document Conversion")
        target_format = st.selectbox(
            "Select target format",
            options=["pdf", "word", "txt", "doc", "image"],
            help="Choose the format you want to convert your document to"
        )
        if st.button("🔄 Convert Document"):
            with st.spinner("Converting document..."):
                converted_file = convert_document(file_path, file_type, target_format)
            if converted_file:
                st.success("Conversion complete!")
                with open(converted_file, "rb") as f:
                    st.download_button(
                        "⬇️ Download Converted Document",
                        data=f,
                        file_name=os.path.basename(converted_file),
                        help="Click to download your converted document"
                    )
            else:
                st.error("Conversion failed. Please check your file and target format.")

# Add footer
st.markdown("---")
st.markdown("""
    <div style='text-align: center'>
        <p>Made with ❤️ using Streamlit</p>
    </div>
""", unsafe_allow_html=True)

# Add this function after your other function definitions
def get_youtube_video_id(url):
    """Extract YouTube video ID from URL"""
    video_id = None
    patterns = [
        r'(?:v=|/v/|youtu\.be/|/embed/)([^&?/]+)',
        r'youtube.com/watch\?v=([^&?/]+)',
        r'youtu.be/([^&?/]+)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            video_id = match.group(1)
            break
    
    return video_id

def get_youtube_transcript(video_id):
    """Get transcript from YouTube video with fallback to other languages"""
    try:
        # First try to get English transcript
        try:
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        except:
            # If English fails, try Hindi (auto-generated) and translate
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id, languages=['hi'])
            
            # If you want to try all available languages:
            if not transcript_list:
                available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
                transcript_list = available_transcripts.find_transcript(['hi']).translate('en').fetch()
        
        # Combine all transcript pieces
        transcript = ' '.join([t['text'] for t in transcript_list])
        return transcript
    
    except Exception as e:
        # Get available languages for better error message
        try:
            available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
            available_langs = [
                f"{t.language_code} ({t.language})"
                for t in available_transcripts
            ]
            return f"Error getting transcript. Available languages are: {', '.join(available_langs)}"
        except:
            return f"Error getting transcript: {str(e)}"

# Add this new section before the footer
with st.container():
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.subheader("🎥 YouTube Video Summarizer")
    
    youtube_url = st.text_input(
        "Enter YouTube Video URL",
        placeholder="https://www.youtube.com/watch?v=...",
        help="Paste the URL of the YouTube video you want to summarize"
    )
    
    # Add language selection
    lang_code = st.selectbox(
        "Select transcript language",
        options=['en', 'hi', 'es', 'fr', 'de', 'ja', 'ko', 'zh-Hans'],
        format_func=lambda x: {
            'en': 'English',
            'hi': 'Hindi',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'ja': 'Japanese',
            'ko': 'Korean',
            'zh-Hans': 'Chinese (Simplified)'
        }[x],
        help="Select the language of the transcript you want to use"
    )
    
    if youtube_url and st.button("📝 Summarize Video"):
        with st.spinner("Processing video..."):
            try:
                # Get video ID
                video_id = get_youtube_video_id(youtube_url)
                if not video_id:
                    st.error("Invalid YouTube URL. Please check the URL and try again.")
                else:
                    # Embed video
                    st.video(youtube_url)
                    
                    # Show available transcripts
                    try:
                        available_transcripts = YouTubeTranscriptApi.list_transcripts(video_id)
                        st.info("Available transcript languages: " + 
                               ", ".join([f"{t.language_code} ({t.language})" 
                                        for t in available_transcripts]))
                    except:
                        pass
                    
                    # Get transcript
                    transcript = get_youtube_transcript(video_id)
                    
                    if isinstance(transcript, str) and not transcript.startswith("Error"):
                        # Get summarizer
                        summarizer = get_summarizer()
                        
                        # Split transcript into chunks if it's too long
                        max_chunk_length = 1024
                        chunks = [transcript[i:i + max_chunk_length] 
                                for i in range(0, len(transcript), max_chunk_length)]
                        
                        # Summarize each chunk
                        summaries = []
                        for chunk in chunks:
                            summary = summarizer(chunk, max_length=150, min_length=30, do_sample=False)
                            summaries.append(summary[0]['summary_text'])
                        
                        # Combine summaries
                        final_summary = " ".join(summaries)
                        
                        # Display results in columns
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("### 📌 Video Summary")
                            st.write(final_summary)
                            st.session_state['video_summary'] = final_summary
                            
                            # Text-to-speech option
                            if st.button("🔊 Listen to Summary", key="video_listen_summary"):
                                try:
                                    with st.spinner("Converting to speech..."):
                                        # Create temporary file for audio
                                        with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as fp:
                                            tts = gTTS(text=st.session_state['video_summary'], lang='en')
                                            tts.save(fp.name)
                                            
                                            # Stop any currently playing audio
                                            if pygame.mixer.music.get_busy():
                                                pygame.mixer.music.stop()
                                            
                                            # Load and play the new audio
                                            pygame.mixer.music.load(fp.name)
                                            pygame.mixer.music.play()
                                            
                                            # Wait for audio to finish
                                            while pygame.mixer.music.get_busy():
                                                pygame.time.Clock().tick(10)
                                                st.empty()  # Keep the stream alive
                                            
                                            # Cleanup
                                            pygame.mixer.music.unload()
                                            os.unlink(fp.name)
                                    st.success("Audio playback completed!")
                                except Exception as e:
                                    st.error(f"Error during audio playback: {str(e)}")
                        
                        with col2:
                            st.markdown("### 📊 Video Visualization")
                            viz_type = st.selectbox(
                                "Select visualization type",
                                options=["mindmap", "flowchart", "graph"],
                                key="video_viz_type",
                                help="Choose how you want to visualize the video content"
                            )
                            
                            if st.button("Generate Video Visualization"):
                                with st.spinner("Generating visualization..."):
                                    try:
                                        diagram_code = generate_mindmap(final_summary, viz_type)
                                        st.markdown("### 📊 Video Content Diagram")
                                        st.markdown(f"```mermaid\n{diagram_code}\n```")
                                        
                                        # Add download button for diagram
                                        st.download_button(
                                            "⬇️ Download Diagram Code",
                                            diagram_code,
                                            file_name=f"video_{viz_type}_diagram.mmd",
                                            mime="text/plain",
                                            key="video_diagram_download"
                                        )
                                        
                                        # Show raw code option
                                        if st.checkbox("Show Raw Diagram Code", key="video_show_code"):
                                            st.code(diagram_code, language="mermaid")
                                    except Exception as e:
                                        st.error(f"Error generating diagram: {str(e)}")
                        
                        # Show full transcript option below both columns
                        if st.checkbox("Show Full Transcript"):
                            st.markdown("### 📜 Full Transcript")
                            st.text_area("", transcript, height=200)
                    else:
                        st.error(transcript)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
    
    st.markdown('</div>', unsafe_allow_html=True)

# Add this function after your other function definitions
def generate_mindmap(text, map_type="mindmap"):
    """
    Generate a Mermaid diagram (mindmap/flowchart) from text
    Args:
        text: The text to generate the diagram from
        map_type: "mindmap", "flowchart", or "graph"
    Returns:
        str: Mermaid diagram code
    """
    try:
        # Create a prompt template for the diagram generation
        if map_type == "mindmap":
            template = """
            Create a mindmap in Mermaid syntax based on the following text. 
            Focus on the main concepts and their relationships.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the mindmap like this:
            mindmap
                Root Topic
                    Subtopic 1
                        Detail A
                        Detail B
                    Subtopic 2
                        Detail C
                        Detail D
            """
        elif map_type == "flowchart":
            template = """
            Create a flowchart in Mermaid syntax based on the following text.
            Focus on the process flow and relationships between concepts.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the flowchart like this:
            graph TD
                A[Start] --> B[Process 1]
                B --> C[Process 2]
                C --> D[End]
            """
        else:  # graph
            template = """
            Create a graph diagram in Mermaid syntax based on the following text.
            Focus on the relationships between concepts.
            Use only alphanumeric characters, spaces, and hyphens in the node text.
            Text: {text}
            
            Format the graph like this:
            graph LR
                A[Concept 1] --- B[Concept 2]
                B --- C[Concept 3]
                C --- D[Concept 4]
            """
        
        prompt = PromptTemplate(
            input_variables=["text"],
            template=template
        )
        
        # Use the existing LLM to generate the diagram
        llm = HuggingFaceHub(repo_id="google/flan-t5-large")
        diagram_code = llm.predict(prompt.format(text=text))
        
        return diagram_code.strip()
    
    except Exception as e:
        return f"Error generating diagram: {str(e)}"

# Create FastAPI app
app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # the Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Mount static files (if needed)
app.mount("/static", StaticFiles(directory="static"), name="static")

# API endpoints
@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Process the file and generate summary
        summary = summarize_document(await file.read(), file.filename)
        return JSONResponse(content={"summary": summary})
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

@app.post("/convert")
async def convert_file(file: UploadFile = File(...), format: str = None):
    try:
        # Process the file conversion
        converted_file = convert_document(await file.read(), file.filename, format)
        return JSONResponse(content={"file": converted_file})
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

@app.post("/youtube")
async def process_youtube(data: dict):
    try:
        # Process YouTube URL
        url = data.get("url")
        summary = summarize_youtube_video(url)
        return JSONResponse(content={"summary": summary})
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"error": str(e)}
        )

# Streamlit UI
def main():
    st.set_page_config(
        page_title="Document Interaction Platform",
        page_icon="📄",
        layout="wide"
    )
    
    # Your existing Streamlit UI code here
    # ...

if __name__ == "__main__":
    # Run both FastAPI and Streamlit
    import threading
    
    def run_fastapi():
        uvicorn.run(app, host="0.0.0.0", port=8000)
    
    # Start FastAPI in a separate thread
    threading.Thread(target=run_fastapi, daemon=True).start()
    
    # Run Streamlit
    main()

